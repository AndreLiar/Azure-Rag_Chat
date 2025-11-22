import uuid
from pathlib import Path
from typing import List, Dict, Any
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, and_

# Document processing imports
import docx
import pandas as pd
import pypdf

from .search_service import SearchService
from db_models import Document, DocumentChunk
from database import get_tenant_id


class DocumentService:
    def __init__(self, db: AsyncSession):
        self.db = db
        self.search_service = SearchService()
        self.supported_extensions = {".pdf", ".docx", ".txt", ".csv"}

    async def process_document(self, file_path: str, filename: str, uploaded_by_id: str = None) -> str:
        """Process a document and add it to both database and search index"""
        file_extension = Path(filename).suffix.lower()
        tenant_id = get_tenant_id()

        if not tenant_id:
            raise ValueError("No tenant context available")

        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")

        try:
            # Create document record in database
            document = Document(
                organization_id=tenant_id,
                filename=filename,
                blob_url=f"file://{file_path}",  # In production, this would be Azure Blob URL
                file_size=Path(file_path).stat().st_size,
                status='processing',
                uploaded_by=uploaded_by_id
            )
            self.db.add(document)
            await self.db.flush()  # Get the document ID

            # Extract text content
            content = self._extract_text(file_path, file_extension)

            # Split into chunks
            chunks = self._split_text(content)

            # Create document chunks in database and prepare for search index
            search_documents = []
            
            for i, chunk_content in enumerate(chunks):
                # Create database chunk
                chunk = DocumentChunk(
                    organization_id=tenant_id,
                    document_id=document.id,
                    chunk_index=i,
                    content=chunk_content
                )
                self.db.add(chunk)
                
                # Prepare for search index
                search_doc = {
                    "id": f"{document.id}_{i}",
                    "content": chunk_content,
                    "title": filename,
                    "source": filename,
                    "document_id": str(document.id),  # Link back to database document
                }
                search_documents.append(search_doc)

            # Add to search index
            await self.search_service.add_documents(search_documents)

            # Update document status to complete
            document.status = 'complete'
            await self.db.commit()

            return str(document.id)
        
        except Exception as e:
            await self.db.rollback()
            # Update document status to failed if it exists
            if 'document' in locals():
                document.status = 'failed'
                await self.db.commit()
            raise e

    def _extract_text(self, file_path: str, file_extension: str) -> str:
        """Extract text from different file types"""
        if file_extension == ".pdf":
            return self._extract_from_pdf(file_path)
        elif file_extension == ".docx":
            return self._extract_from_docx(file_path)
        elif file_extension == ".txt":
            return self._extract_from_txt(file_path)
        elif file_extension == ".csv":
            return self._extract_from_csv(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        with open(file_path, "rb") as file:
            reader = pypdf.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT"""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()

    def _extract_from_csv(self, file_path: str) -> str:
        """Extract text from CSV"""
        df = pd.read_csv(file_path)
        return df.to_string()

    def _split_text(
        self, text: str, chunk_size: int = 1000, overlap: int = 200
    ) -> List[str]:
        """Split text into overlapping chunks"""
        words = text.split()
        chunks = []

        for i in range(0, len(words), chunk_size - overlap):
            end_idx = i + chunk_size
            chunk = " ".join(words[i:end_idx])
            if chunk.strip():  # Only add non-empty chunks
                chunks.append(chunk)

        return chunks

    async def list_documents(self) -> List[Dict[str, Any]]:
        """List all documents for the current tenant"""
        tenant_id = get_tenant_id()
        if not tenant_id:
            raise ValueError("No tenant context available")

        result = await self.db.execute(
            select(Document).where(Document.organization_id == tenant_id)
        )
        documents = result.scalars().all()
        
        return [
            {
                "id": str(doc.id),
                "title": doc.filename,
                "source": doc.filename,
                "status": doc.status,
                "file_size": doc.file_size,
                "created_at": doc.created_at.isoformat() if doc.created_at else None
            }
            for doc in documents
        ]

    async def delete_document(self, document_id: str):
        """Delete a document and all its chunks from both database and search index"""
        tenant_id = get_tenant_id()
        if not tenant_id:
            raise ValueError("No tenant context available")

        try:
            # Get document from database (with tenant check via RLS)
            result = await self.db.execute(
                select(Document).where(
                    and_(
                        Document.id == document_id,
                        Document.organization_id == tenant_id
                    )
                )
            )
            document = result.scalar_one_or_none()
            
            if not document:
                raise ValueError(f"Document {document_id} not found")

            # Get all chunks for this document from database
            chunk_result = await self.db.execute(
                select(DocumentChunk).where(DocumentChunk.document_id == document_id)
            )
            chunks = chunk_result.scalars().all()

            # Delete from search index first
            for chunk in chunks:
                search_chunk_id = f"{document_id}_{chunk.chunk_index}"
                try:
                    await self.search_service.delete_document(search_chunk_id)
                except Exception as e:
                    print(f"Warning: Failed to delete chunk {search_chunk_id} from search index: {e}")

            # Delete from database (chunks will be deleted via cascade)
            await self.db.delete(document)
            await self.db.commit()

        except Exception as e:
            await self.db.rollback()
            raise e
